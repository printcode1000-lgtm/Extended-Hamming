"""Build editable Word versions of the public academic PDF library."""

from __future__ import annotations

import csv
import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
AUTHOR = "Howida Gharib Saad El Din Selim"
VERSION = "Research Project Version 0.2 - Validated Draft"
TODAY = date.today().isoformat()


summary = json.loads((ROOT / "thesis/data/summary.json").read_text(encoding="utf-8"))
functional = json.loads((ROOT / "results/functional_verification.json").read_text(encoding="utf-8"))
nominal = {f"{r['level']}:{r['architecture']}": r for r in summary["nominal"]}

REFERENCES = [
    "R. W. Hamming, 'Error Detecting and Error Correcting Codes,' Bell System Technical Journal, vol. 29, no. 2, pp. 147-160, 1950.",
    "A. Morgenshtein, A. Fish, and I. A. Wagner, 'Gate-Diffusion Input (GDI): A Power-Efficient Method for Digital Combinatorial Circuits,' IEEE TVLSI, 2002.",
    "A. Morgenshtein et al., 'Full-Swing Gate Diffusion Input Logic - Case-Study of Low-Power CLA Adder Design,' Integration, 2014.",
    "M. A. M. El-Bendary and O. El-Badry, 'FS-GDI Based Area Efficient Hamming (11,7) Encoding,' International Journal of Electronics, 2024.",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 22, "183B5B"),
        ("Subtitle", 12, "4E5964"),
        ("Heading 1", 16, "183B5B"),
        ("Heading 2", 13, "183B5B"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if name.startswith("Heading"):
            style.font.bold = True
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(5)


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = f"{AUTHOR} | {VERSION}"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(93, 104, 115)


def add_cover(doc: Document, title: str, label: str) -> None:
    doc.add_paragraph(label.upper(), style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Researcher and Author\n{AUTHOR}", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(VERSION, style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(TODAY, style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def p(doc: Document, text: str, *, italic: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = italic


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def table(doc: Document, rows: list[list[object]]) -> None:
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    tbl.autofit = True
    set_repeat_table_header(tbl.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = str(value)
            if r_idx == 0:
                set_cell_shading(cell, "DDE6ED")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(24, 59, 91)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8.5)
    doc.add_paragraph()


def markdown(doc: Document, rel_path: str, *, skip_title: bool = False) -> None:
    lines = (ROOT / rel_path).read_text(encoding="utf-8").splitlines()
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            p(doc, " ".join(paragraph).replace("**", "").replace("*", ""))
            paragraph.clear()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if skip_title and i == 0 and line.startswith("# "):
            i += 1
            continue
        if line.startswith("# "):
            flush()
            h1(doc, line[2:])
            i += 1
            continue
        if line.startswith("## "):
            flush()
            h2(doc, line[3:])
            i += 1
            continue
        if line.startswith("### "):
            flush()
            h2(doc, line[4:])
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|---"):
            flush()
            block = [line]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = [[cell.strip().strip("`") for cell in row.strip("|").split("|")] for row in block]
            table(doc, rows)
            continue
        if line.startswith("- "):
            flush()
            bullet(doc, re.sub(r"^- ", "", line).replace("**", ""))
            i += 1
            continue
        if re.match(r"\d+\.\s", line):
            flush()
            numbered(doc, re.sub(r"^\d+\.\s", "", line).replace("**", ""))
            i += 1
            continue
        if not line:
            flush()
            i += 1
            continue
        paragraph.append(line)
        i += 1
    flush()


def nominal_table(doc: Document, level: str = "EH127") -> None:
    rows = [["Architecture", "Power (uW)", "Delay (ps)", "PDP (fJ)", "Swing (V)", "Transistors", "Status"]]
    for r in summary["nominal"]:
        if r["level"] == level:
            rows.append([
                r["architecture"],
                f"{r['pavg'] * 1e6:.4f}",
                f"{r['tpd'] * 1e12:.3f}",
                f"{r['pdp'] * 1e15:.4f}",
                f"{r['output_swing']:.4f}",
                r["transistor_count"],
                "Pass" if r["functional_pass"] else "Fail",
            ])
    table(doc, rows)


def architecture_summary(doc: Document) -> None:
    h1(doc, "Hybrid-B Architecture Summary")
    bullet(doc, "Data inputs D1 through D7 feed FS-GDI XOR parity networks.")
    bullet(doc, "The overall parity path adds P0 for SEC-DED detection.")
    bullet(doc, "CMOS restoration is applied only to externally observed parity outputs.")
    bullet(doc, "Hybrid-B is a speed-oriented restoration compromise, not a universal energy optimum.")


def add_figures(doc: Document) -> None:
    for name, caption in [
        ("eh127_nominal_power.png", "Figure 8.1. Nominal average power."),
        ("eh127_nominal_delay.png", "Figure 8.2. Nominal D1-to-P0 delay."),
        ("eh127_nominal_pdp.png", "Figure 8.3. Nominal power-delay product."),
        ("eh127_pdp_vs_vdd.png", "Figure 8.4. PDP versus supply voltage."),
    ]:
        path = ROOT / "figures" / name
        if path.exists():
            doc.add_picture(str(path), width=Inches(6.1))
            para = doc.add_paragraph(caption)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(78, 89, 100)


def add_references(doc: Document) -> None:
    h1(doc, "References")
    for idx, ref in enumerate(REFERENCES, 1):
        p(doc, f"[{idx}] {ref}")


def thesis_content(doc: Document) -> None:
    h1(doc, "Abstract")
    p(doc, "This working thesis draft evaluates CMOS, GDI, FS-GDI, and hybrid FS-GDI/CMOS implementations of Hamming (11,7) and extended Hamming (12,7). The baseline paper is independently audited, SEC-DED logic is exhaustively verified, and 121 LTspice experiments are executed with a predictive 65-nm PTM model. Of these, 118 measurement sets pass automated validation and three low-load cases are rejected for excessive overshoot. Hybrid-B provides the lowest measured D1-to-P0 delay at the nominal condition, whereas FS-GDI retains the lowest simulated power and PDP.")
    h1(doc, "Abbreviations")
    table(doc, [["Abbreviation", "Meaning"], ["CMOS", "Complementary metal-oxide-semiconductor"], ["GDI", "Gate diffusion input"], ["FS-GDI", "Full-swing gate diffusion input"], ["PDP", "Power-delay product"], ["PTM", "Predictive technology model"], ["SEC-DED", "Single-error correction, double-error detection"]])
    h1(doc, "Chapter 1 - Introduction")
    p(doc, "Low-power error-control hardware is dominated by XOR networks. Device count alone does not establish useful operation because pass-transistor degradation, output load, restoration overhead, and cascade depth can change the power-delay trade-off.")
    h1(doc, "Chapter 2 - Literature Survey")
    p(doc, "Hamming established the coding theory; Morgenshtein and co-authors established GDI and later FS-GDI. Recent ECC work addresses SEC-DED and stronger adjacent-error codes. Extended Hamming and GDI are therefore not new by themselves.")
    h1(doc, "Chapter 3 - Theoretical Background")
    p(doc, "P1 = D1 XOR D2 XOR D4 XOR D5 XOR D7")
    p(doc, "P2 = D1 XOR D3 XOR D4 XOR D6 XOR D7")
    p(doc, "P4 = D2 XOR D3 XOR D4; P8 = D5 XOR D6 XOR D7")
    p(doc, "P0 = D1 XOR ... XOR D7 XOR P1 XOR P2 XOR P4 XOR P8")
    h1(doc, "Chapter 4 - Baseline Reproduction")
    nominal_table(doc, "H117")
    h1(doc, "Chapter 5 - Extended Hamming (12,7)")
    p(doc, "The extended architecture adds a ten-XOR overall-parity path to the twelve-XOR Hamming network.")
    architecture_summary(doc)
    h1(doc, "Chapter 7 - LTspice Simulation Methodology")
    table(doc, [["Parameter", "Value"], ["Simulator", "LTspice 26.0.2"], ["Model", "PTM 65-nm bulk CMOS beta"], ["Nominal", "1.2 V, 27 C, 125 MHz, 10 fF"], ["Sizing", "NMOS 120/65 nm; PMOS 240/65 nm"], ["Delay", "50%-to-50%; D1-to-P0 for EH127"], ["Power", "AVG(-V(VDD)*I(VSUP))"], ["Sweeps", "VDD, temperature, load, and frequency independently"]])
    h1(doc, "Chapter 8 - Results and Robustness Analysis")
    p(doc, f"Exhaustive verification covered {functional['data_words_tested']} data words, {functional['single_bit_error_cases_tested']} single-bit cases, and {functional['double_bit_error_cases_tested']} double-bit cases. All passed.")
    nominal_table(doc, "EH127")
    add_figures(doc)
    hb = nominal["EH127:HYBRID_B"]
    cm = nominal["EH127:CMOS"]
    fs = nominal["EH127:FSGDI"]
    h1(doc, "Chapter 9 - Discussion")
    p(doc, f"At nominal conditions, Hybrid-B uses {hb['pavg'] * 1e6:.3f} uW and has a D1-to-P0 delay of {hb['tpd'] * 1e12:.3f} ps. Relative to CMOS, this is {((cm['pavg'] - hb['pavg']) / cm['pavg'] * 100):.2f}% lower power and {((cm['tpd'] - hb['tpd']) / cm['tpd'] * 100):.2f}% lower delay. FS-GDI uses only {fs['pavg'] * 1e6:.3f} uW but has a longer P0 path.")
    h1(doc, "Chapter 10 - Conclusions and Future Work")
    p(doc, "Hybrid-B is supported as a speed-oriented selective-restoration candidate, not as a universal energy optimum. Future work requires supervisor review, an independent foundry-PDK or second-simulator cross-check, post-layout parasitics, and a lower-voltage boundary search.")
    add_references(doc)


def literature_content(doc: Document) -> None:
    h1(doc, "Survey Methodology")
    p(doc, "The survey prioritizes 2021-2026 publications while retaining foundational Hamming, GDI, and FS-GDI sources. Publisher pages and DOI records are used to verify existence. Unavailable metadata are marked rather than inferred.")
    rows = list(csv.DictReader((ROOT / "literature/survey/survey.csv").open(encoding="utf-8")))
    compact = [["Reference", "Year", "Circuit", "Logic", "Contribution", "Limitation"]]
    for r in rows:
        compact.append([r["Reference"], r["Year"], r["Circuit/Code"], r["Logic Style"], r["Main Contribution"], r["Limitation"]])
    table(doc, compact)
    markdown(doc, "docs/research-gap.md", skip_title=True)


def results_content(doc: Document) -> None:
    markdown(doc, "docs/simulation-methodology.md", skip_title=True)
    h1(doc, "Nominal Comparison")
    nominal_table(doc, "EH127")
    hb = nominal["EH127:HYBRID_B"]
    cm = nominal["EH127:CMOS"]
    fs = nominal["EH127:FSGDI"]
    p(doc, f"Hybrid-B reduces nominal power by {(cm['pavg'] - hb['pavg']) / cm['pavg'] * 100:.2f}%, delay by {(cm['tpd'] - hb['tpd']) / cm['tpd'] * 100:.2f}%, PDP by {(cm['pdp'] - hb['pdp']) / cm['pdp'] * 100:.2f}%, and transistor count by {(cm['transistor_count'] - hb['transistor_count']) / cm['transistor_count'] * 100:.2f}% relative to CMOS.")
    p(doc, f"FS-GDI has the lowest nominal power ({fs['pavg'] * 1e6:.3f} uW) and PDP ({fs['pdp'] * 1e15:.3f} fJ), but its P0 path is slower than CMOS.")
    add_figures(doc)


def simple_content(doc: Document, sources: list[str], additions=None) -> None:
    for source in sources:
        markdown(doc, source, skip_title=False)
    if additions:
        additions(doc)


def proposed_design(doc: Document) -> None:
    h1(doc, "Code Architecture")
    p(doc, "The design adds P0 to Hamming (11,7) and evaluates two restoration placements.")
    p(doc, "P0 = D1 XOR ... XOR D7 XOR P1 XOR P2 XOR P4 XOR P8")
    architecture_summary(doc)
    h1(doc, "Circuit-Level Alternatives")
    p(doc, "CMOS uses 12-transistor XOR cells. The reconstructed FS-GDI cell uses six transistors. Hybrid-A adds local CMOS restoration to every cell. Hybrid-B adds non-inverting restoration only at parity outputs.")
    h1(doc, "Measured Trade-off")
    nominal_table(doc, "EH127")
    p(doc, "Hybrid-B is selected for speed; FS-GDI remains the energy optimum under the tested nominal condition.")


def software_tools(doc: Document) -> None:
    h1(doc, "Tools Actually Used")
    table(doc, [["Tool", "Version", "Role", "Official source"], ["LTspice", "26.0.2", "Transistor-level simulation", "analog.com/ltspice"], ["Python", "3.x runtime", "Verification, parsing, PDF, Word, and figure generation", "python.org"], ["Git", "Installed client", "Version control", "git-scm.com"], ["Icarus Verilog", "12.x", "SystemVerilog reference simulation", "bleyer.org/icarus"], ["GitHub", "Public repository service", "Public dissemination", "github.com"], ["PTM", "65-nm beta model", "Predictive transistor model", "mec.umn.edu/ptm"], ["ReportLab", "Installed Python library", "Academic PDF generation", "reportlab.com"], ["python-docx", "Installed Python library", "Editable Word document generation", "python-docx.readthedocs.io"], ["Pillow", "Installed Python library", "Research figures", "python-pillow.org"]])
    h1(doc, "Reproducibility Notes")
    p(doc, "Versions and commands are documented in the repository. Verilog source is included and can be simulated with Icarus Verilog for the reference encoder testbench.")


def repository_guide(doc: Document) -> None:
    h1(doc, "Public Repository")
    p(doc, "https://github.com/printcode1000-lgtm/Extended-Hamming")
    h1(doc, "Access")
    p(doc, "Clone with: git clone https://github.com/printcode1000-lgtm/Extended-Hamming.git. Use the main branch. LTspice files are under ltspice/, scripts under scripts/, validated data under data/processed/, PDFs and Word files under pdf/, and the website at the repository root.")
    h1(doc, "Citation Status")
    p(doc, "No DOI has been assigned. Cite the researcher, repository title, URL, version, and access date until a formal archival identifier exists.")


def reproducibility_guide(doc: Document) -> None:
    h1(doc, "Procedure")
    for step in [
        "Clone the main branch.",
        "Install LTspice 26.0.2 and Python 3.",
        "Confirm models/65nm_bulk.pm.",
        "Run scripts/verify_secded.py.",
        "Run scripts/build_ltspice.py.",
        "Run scripts/run_ltspice.py with the LTspice path.",
        "Run results_pipeline.py and summarize_results.py.",
        "Run generate_figures.py, generate_academic_pdfs.py, and generate_academic_word.py.",
        "Run link_check.py and the secret scan.",
        "Open the website through a local server under /Extended-Hamming/.",
    ]:
        numbered(doc, step)
    h1(doc, "Validation")
    p(doc, "A result is accepted only if required .meas values exist, PDP recalculates correctly, and signal bounds pass the configured checks. Rejected raw logs remain retained.")


def secded_verification(doc: Document) -> None:
    h1(doc, "Reference Model")
    p(doc, "The Python model uses positions 1, 2, 4, and 8 for Hamming parity and position 12 for overall even parity.")
    h1(doc, "Exhaustive Results")
    table(doc, [["Test", "Cases", "Result"], ["Data words", functional["data_words_tested"], "Pass"], ["Single-bit errors", functional["single_bit_error_cases_tested"], "Corrected"], ["Double-bit errors", functional["double_bit_error_cases_tested"], "Detected"]])
    h1(doc, "Scope")
    p(doc, "This verifies encoder/decoder logic. It does not claim a transistor-level decoder implementation.")


def technical_architecture(doc: Document) -> None:
    h1(doc, "Research Flow")
    p(doc, "Verified literature -> audited equations -> Python and Verilog references -> generated LTspice netlists -> raw logs -> validated CSV/JSON -> figures -> website, PDFs, and Word documents.")
    h1(doc, "Repository Layers")
    table(doc, [["Layer", "Location", "Purpose"], ["Theory and audit", "literature/, docs/", "Sources, gap, and decisions"], ["Simulation", "ltspice/, models/", "Circuit and device sources"], ["Verification", "scripts/, verification/", "Functional and data validation"], ["Evidence", "data/, results/, figures/", "Raw-to-processed provenance"], ["Publication", "thesis/, pdf/", "Public academic outputs"]])
    h1(doc, "Provenance")
    p(doc, "Each measurement is connected through an experiment ID to a circuit, model, condition set, raw log, processed row, and final table or figure.")


def build(pdf_path: Path, title: str, label: str, content_fn) -> Path:
    docx_path = pdf_path.with_suffix(".docx")
    doc = Document()
    style_doc(doc)
    add_footer(doc)
    doc.core_properties.author = AUTHOR
    doc.core_properties.title = title
    doc.core_properties.subject = "Extended Hamming (12,7), FS-GDI, CMOS, Low-Power VLSI"
    doc.core_properties.keywords = "Extended Hamming, SEC-DED, FS-GDI, CMOS, LTspice, 65 nm"
    add_cover(doc, title, label)
    content_fn(doc)
    doc.save(docx_path)
    return docx_path


def main() -> None:
    docs = [
        (ROOT / "pdf/thesis/Howida_Gharib_Extended_Hamming_MSc_Thesis.pdf", "Design and Robustness Analysis of a Low-Power Hybrid FS-GDI/CMOS Extended Hamming (12,7) Encoder", "Working Thesis Draft", thesis_content),
        (ROOT / "pdf/research/Literature_Survey.pdf", "Literature Survey", "Research Document", literature_content),
        (ROOT / "pdf/research/Research_Gap_and_Contribution.pdf", "Research Gap and Contribution", "Research Document", lambda d: simple_content(d, ["docs/research-gap.md"])),
        (ROOT / "pdf/research/Original_Paper_Audit.pdf", "Original Paper Audit", "Research Audit", lambda d: simple_content(d, ["docs/original-paper-audit.md"])),
        (ROOT / "pdf/research/Proposed_Hybrid_Extended_Hamming_Encoder.pdf", "Proposed Hybrid Extended Hamming (12,7) Encoder", "Design Document", proposed_design),
        (ROOT / "pdf/technical/LTspice_Simulation_Methodology.pdf", "LTspice Simulation Methodology", "Technical Document", lambda d: simple_content(d, ["docs/simulation-methodology.md", "docs/measurement-definitions.md"])),
        (ROOT / "pdf/technical/65nm_Device_Model_Documentation.pdf", "65-nm Device Model Documentation", "Technical Document", lambda d: simple_content(d, ["docs/model-documentation.md"])),
        (ROOT / "pdf/technical/Simulation_Measurement_Definitions.pdf", "Simulation Measurement Definitions", "Technical Document", lambda d: simple_content(d, ["docs/measurement-definitions.md"])),
        (ROOT / "pdf/technical/Software_and_Research_Tools.pdf", "Software and Research Tools", "Technical Document", software_tools),
        (ROOT / "pdf/reproducibility/GitHub_Public_Repository_Guide.pdf", "GitHub Public Repository Guide", "Repository Document", repository_guide),
        (ROOT / "pdf/reproducibility/Reproducibility_Guide.pdf", "Reproducibility Guide", "Research Reproducibility", reproducibility_guide),
        (ROOT / "pdf/simulation/Simulation_Results_and_Analysis.pdf", "Simulation Results and Analysis", "Simulation Report", results_content),
        (ROOT / "pdf/simulation/SEC-DED_Verification_Documentation.pdf", "SEC-DED Verification Documentation", "Verification Report", secded_verification),
        (ROOT / "pdf/technical/Project_Technical_Architecture.pdf", "Project Technical Architecture", "Technical Document", technical_architecture),
        (ROOT / "pdf/publication/Publication_Plan_and_Manuscript_Outline.pdf", "Publication Plan and Manuscript Outline", "Publication Material", lambda d: simple_content(d, ["docs/publication-plan.md", "docs/paper-outline.md"])),
    ]
    for pdf_path, title, label, content_fn in docs:
        out = build(pdf_path, title, label, content_fn)
        print(out.relative_to(ROOT))


if __name__ == "__main__":
    main()
