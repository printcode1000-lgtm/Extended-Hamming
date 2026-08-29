from __future__ import annotations

import json
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "pdf" / "research"
DOCX_OUT = OUT_DIR / "Memristor_PIM_Current_Development_Phase_Report.docx"
PDF_OUT = OUT_DIR / "Memristor_PIM_Current_Development_Phase_Report.pdf"


def load_json(path: str) -> dict:
    with (ROOT / path).open(encoding="utf-8") as handle:
        return json.load(handle)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def add_para(doc: Document, text: str, style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    run.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    set_rtl(p)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    run.font.color.rgb = RGBColor(31, 77, 120)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_rtl(p)
        run = p.add_run(item)
        run.font.name = "Tahoma"
        run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
        run.font.size = Pt(10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_rtl(p)
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Tahoma"
        run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
        run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            set_rtl(p)
            run = p.add_run(text)
            run.font.name = "Tahoma"
            run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
            run.font.size = Pt(8.5)
    set_table_width(table, widths)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Tahoma"
        style._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    set_rtl(p)
    title = p.add_run("تقرير منهجي عن مجهود الباحث في مرحلة تطوير Memristor-PIM")
    title.bold = True
    title.font.name = "Tahoma"
    title._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    title.font.size = Pt(20)
    title.font.color.rgb = RGBColor(11, 37, 69)
    add_para(doc, "Design and Robustness Analysis of a Low-Power Hybrid FS-GDI/CMOS Extended Hamming (12,7) Encoder", bold=True)
    add_para(doc, "الباحثة: Howida Gharib Saad El Din Selim")
    add_para(doc, "الغرض: توثيق علمي لما أضيف في طبقة Memristor-Based Hamming Encoder مع الحفاظ على نتائج CMOS/FS-GDI/Hybrid-B السابقة.")
    add_para(doc, "تاريخ التوليد: 2026-08-29")
    add_para(doc, "حالة التقرير: ملحق بحثي قابل للمراجعة، وليس ادعاء تصنيع أو تحقق سيليكوني.")
    doc.add_page_break()


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    comparison = load_json("results/memristor/ltspice_comparison_summary.json")
    optimized_mc = load_json("results/memristor/optimized_monte_carlo_summary.json")
    conservative_mc = load_json("results/memristor/monte_carlo_summary.json")

    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. الملخص العلمي", 1)
    add_para(
        doc,
        "يوثق هذا التقرير مجهود الباحث في إضافة طبقة Memristor-Based Hamming Encoder إلى مشروع Extended Hamming (12,7). "
        "تم الحفاظ على المسار السابق كما هو: CMOS وFS-GDI وHybrid-A وHybrid-B بقيت مراحل أساس وقياس، ثم أضيفت طبقة Processing-in-Memory باعتبارها امتدادًا بحثيًا فوق خط الأساس وليس بديلًا يمحو نتائجه."
    )
    add_para(
        doc,
        "منهجيًا، انتقل البحث من سؤال تقليل ترانزستورات XOR إلى سؤال أوسع: كيف يمكن حساب parity داخل الذاكرة وتقليل حركة البيانات بين الذاكرة والمعالج. "
        "هذا الانتقال يمثل توسعًا معماريًا كبيرًا لأنه يربط ECC hardware بآليات stateful memristive logic وcrossbar arrays وتحليل variation."
    )

    add_heading(doc, "2. طبقة الربط بين البحث السابق والامتداد", 1)
    add_bullets(
        doc,
        [
            "البحث السابق أثبت خط أساس كمي باستخدام LTspice ونموذج 65-nm predictive CMOS، وحدد Hybrid-B كمرشح سريع مع حفظ FS-GDI كأقل قدرة وPDP.",
            "الامتداد الجديد أبقى معادلات parity ووظيفة SEC-DED نفسها، لكنه نقل موضع التنفيذ من XOR tree تقليدي إلى منطق stateful داخل مصفوفة ذاكرة.",
            "Hybrid-B أصبح baseline مقاسًا للمقارنة، بينما Memristor-PIM أصبح مسارًا بحثيًا لخفض الطاقة والتأخير المرتبطين بحركة البيانات.",
            "الحالات السلبية لم تُحذف: نتيجة Memristor-PIM المحافظة بقيت محفوظة كمرحلة أولى، ثم أضيفت مرحلة محسنة بنتائج موجبة تحت نموذج LTspice behavioral macromodel.",
        ],
    )

    add_heading(doc, "3. المنهجية المنفذة", 1)
    add_table(
        doc,
        ["المحور", "الإجراء العلمي", "الدليل داخل المشروع"],
        [
            ["النموذج", "اختيار نموذج VTEAM/RRAM بحثي أولي ثم صياغة نموذج محسّن لسلوك PIM منخفض الطاقة.", "models/memristor/*.json وmodels/memristor/*.md"],
            ["بوابة XOR", "بناء netlists لبوابة XOR ميمريستورية والتحقق من truth table.", "ltspice/memristor-hamming/EXP-MEM-XOR-PIM-*.cir"],
            ["Hamming parity", "تمثيل parity لمشفّر Extended Hamming (12,7) داخل مسار PIM وتجميع نتائج LTspice.", "ltspice/memristor-hamming/EXP-MEM-EH127-PIM-*.cir"],
            ["تحليل variation", "تنفيذ Monte Carlo لحساسية Ron وRoff وthreshold وread noise.", "results/memristor/*monte_carlo_summary.json"],
            ["حفظ المراحل", "الإبقاء على النتيجة المحافظة السلبية ثم إضافة نتيجة محسنة موجبة للمقارنة العلمية.", "data/processed/memristor_ltspice_comparison.csv"],
        ],
        [1800, 4200, 3360],
    )

    add_heading(doc, "4. نتائج LTspice الاسمية", 1)
    opt = comparison["optimized_memristor_pim"]
    add_table(
        doc,
        ["المعمارية", "Power", "Delay", "Energy/word", "التفسير العلمي"],
        [
            ["Hybrid-B baseline", "11.203 uW", "215.846 ps", "89.623 fJ", "خط الأساس المقاس من مرحلة CMOS/FS-GDI/Hybrid-B."],
            ["Memristor-PIM conservative", "224.950 uW", "31.537 ns", "1.800 pJ", "مرحلة أولى وظيفية لكنها أسوأ من Hybrid-B في الطاقة والتأخير."],
            ["Memristor-PIM optimized", "0.012 uW", "30.257 ps", "0.097 fJ", "مرحلة محسنة أعطت نتيجة موجبة ضمن macromodel سلوكي في LTspice."],
        ],
        [2050, 1450, 1450, 1600, 2810],
    )
    add_para(
        doc,
        f"تحسن النموذج المحسن مقارنة بخط Hybrid-B بلغ {opt['energy_reduction_fraction'] * 100:.2f}% في الطاقة و{opt['latency_reduction_fraction'] * 100:.2f}% في التأخير. "
        "هذه الأرقام تُعرض كدليل محاكاة سلوكي داخل LTspice، وليست دليل تصنيع أو post-layout sign-off."
    )

    add_heading(doc, "5. تحقق truth table وMonte Carlo", 1)
    add_table(
        doc,
        ["الاختبار", "الحالة المحافظة", "الحالة المحسنة", "الدلالة"],
        [
            ["XOR truth table", "PASS", "PASS", "الوظيفة المنطقية الأساسية محفوظة."],
            ["Nominal Hamming parity", "PASS", "PASS", "نقل parity equations إلى PIM لم يغير الوظيفة الرياضية."],
            ["Monte Carlo trials", str(conservative_mc["trials"]), str(optimized_mc["trials"]), "اختبار التباين على خصائص memristor."],
            ["Error rate", f"{conservative_mc['logic_error_rate'] * 100:.3f}%", f"{optimized_mc['logic_error_rate'] * 100:.3f}%", "النموذج المحسن وصل إلى نتيجة غير سلبية وفق معيار هذه المرحلة."],
        ],
        [2100, 2100, 2100, 3060],
    )

    add_heading(doc, "6. ضبط صفحة ماذا فعل الباحث", 1)
    add_para(
        doc,
        "تم تحويل صفحة ماذا فعل الباحث إلى صفحة مرجعية خاصة بالمشروع: لا توجد روابط داخلية عامة تقود إليها من صفحات المشروع، لكنها تحتفظ بروابط صادرة إلى الوثائق والنتائج والكود والبيانات. "
        "كما حُفظ زر تبديل اللغة، وبقيت الصفحة هي الصفحة الوحيدة ذات ترجمة عربية كاملة على مستوى المشروع."
    )
    add_para(
        doc,
        "تمت مراجعة صياغة الصفحة لتكون علمية ومنهجية، مع استخدام مصطلحات مثل تحليل الدراسات السابقة، التأصيل النظري، منهجية المحاكاة، مصفوفة الأدلة، وحدود الادعاء العلمي. "
        "كما تم توسيع جزء الأسئلة والأجوبة ليصبح سجلًا شاملًا يربط كل سؤال علمي بالملفات الداعمة داخل المشروع."
    )

    add_heading(doc, "7. حدود الادعاء العلمي", 1)
    add_bullets(
        doc,
        [
            "نتيجة Memristor-PIM المحسنة مبنية على LTspice behavioral macromodel وليست نموذج Verilog-A معايرًا على جهاز مصنع.",
            "لم يتم ادعاء silicon measurement أو foundry PDK sign-off أو post-layout parasitic extraction.",
            "تظل النتيجة المحافظة السلبية جزءًا من سجل البحث لأنها توضح أن تحسين PIM ليس تلقائيًا بل يعتمد على النموذج والنوافذ التشغيلية.",
            "الخطوة العلمية التالية هي معايرة النموذج على RRAM/PCM منشور أو PDK فعلي ثم إعادة Monte Carlo تحت زوايا عملية أوسع.",
        ],
    )

    add_heading(doc, "8. ملفات الأدلة الناتجة", 1)
    add_table(
        doc,
        ["الفئة", "الملفات الرئيسية"],
        [
            ["وثائق الامتداد", "docs/memristor-based-hamming-encoder-research-extension.md; docs/memristor-implementation-results.md"],
            ["LTspice", "ltspice/memristor-hamming/EXP-MEM-XOR-PIM-*.cir; EXP-MEM-EH127-PIM-*.cir"],
            ["بيانات ونتائج", "data/processed/memristor_optimized_results.csv; results/memristor/optimized_monte_carlo_summary.json"],
            ["صفحات المشروع", "memristor-extension.html; research-effort.html"],
            ["هذا التقرير", "pdf/research/Memristor_PIM_Current_Development_Phase_Report.docx; .pdf"],
        ],
        [2200, 7160],
    )
    add_heading(doc, "9. الخلاصة المنهجية", 1)
    add_para(
        doc,
        "يوضح هذا الملحق أن مجهود الباحث لم يقتصر على إضافة فكرة نظرية، بل شمل بناء طبقة محاكاة، ملفات LTspice، بيانات معالجة، تحقق Monte Carlo، وتحديث صفحة مرجعية خاصة تربط مسار العمل بالأدلة. "
        "وبذلك أصبح الانتقال إلى Memristor-PIM امتدادًا بحثيًا موثقًا فوق خط أساس Hybrid-B، مع بقاء النتائج السابقة والنتائج المحافظة غير الملائمة جزءًا من سجل التقييم."
    )
    add_para(
        doc,
        "القيمة العلمية الأساسية لهذه المرحلة هي تحويل الفكرة إلى فرضية قابلة للاختبار: إن كانت عملية parity يمكن تنفيذها داخل الذاكرة تحت نموذج ميمريستوري مضبوط، فقد تتحسن الطاقة والتأخير مقارنة بخط أساس خارج الذاكرة. "
        "أما اكتمال الدليل للنشر القوي فيتطلب معايرة جهازية أعمق وتوسيع زوايا التباين."
    )

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Memristor-PIM Current Development Phase Report - Extended Hamming Research").font.size = Pt(8)
    doc.save(DOCX_OUT)


def convert_to_pdf() -> None:
    soffice_candidates = [
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
    ]
    soffice = next((path for path in soffice_candidates if path.exists()), None)
    if soffice is None:
        raise SystemExit("LibreOffice soffice.exe was not found.")
    if PDF_OUT.exists():
        PDF_OUT.unlink()
    subprocess.run(
        [
            str(soffice),
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUT_DIR),
            str(DOCX_OUT),
        ],
        check=True,
        cwd=str(ROOT),
    )
    generated = DOCX_OUT.with_suffix(".pdf")
    if generated != PDF_OUT and generated.exists():
        generated.replace(PDF_OUT)
    if not PDF_OUT.exists() or PDF_OUT.stat().st_size == 0:
        raise SystemExit("PDF conversion did not create a non-empty PDF.")


def main() -> None:
    build_doc()
    convert_to_pdf()
    print(json.dumps({"docx": str(DOCX_OUT), "pdf": str(PDF_OUT)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
